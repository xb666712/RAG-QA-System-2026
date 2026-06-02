import os
import logging
from typing import List, Tuple, Optional
from PyPDF2 import PdfReader
from docx import Document

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

def read_pdf(file_path: str) -> str:
    """读取PDF文件内容"""
    try:
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
            logger.info(f"成功读取PDF文件: {file_path}, 页数: {len(reader.pages)}")
            return text
    except Exception as e:
        logger.error(f"读取PDF文件失败 {file_path}: {str(e)}")
        return ""

def read_docx(file_path: str) -> str:
    """读取DOCX文件内容"""
    try:
        doc = Document(file_path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        logger.info(f"成功读取DOCX文件: {file_path}, 段落数: {len(doc.paragraphs)}")
        return text
    except Exception as e:
        logger.error(f"读取DOCX文件失败 {file_path}: {str(e)}")
        return ""

def read_txt(file_path: str) -> str:
    """读取TXT文件内容"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
            logger.info(f"成功读取TXT文件: {file_path}, 字符数: {len(text)}")
            return text
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="gbk") as f:
                text = f.read()
                logger.info(f"成功读取TXT文件(GBK编码): {file_path}, 字符数: {len(text)}")
                return text
        except Exception as e:
            logger.error(f"读取TXT文件失败 {file_path}: {str(e)}")
            return ""
    except Exception as e:
        logger.error(f"读取TXT文件失败 {file_path}: {str(e)}")
        return ""

def load_documents_from_folder(folder_path: str) -> List[Tuple[str, str]]:
    """
    从文件夹加载所有支持的文档
    
    Args:
        folder_path: 文件夹路径
        
    Returns:
        文档列表，每个元素为 (文件名, 文件内容) 元组
    """
    documents = []
    supported_extensions = (".pdf", ".docx", ".txt")
    
    if not os.path.exists(folder_path):
        logger.warning(f"文件夹不存在: {folder_path}")
        return documents
    
    if not os.path.isdir(folder_path):
        logger.warning(f"路径不是文件夹: {folder_path}")
        return documents
    
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        
        if not os.path.isfile(file_path):
            continue
        
        lower_filename = filename.lower()
        text = ""
        
        if lower_filename.endswith(".pdf"):
            text = read_pdf(file_path)
        elif lower_filename.endswith(".docx"):
            text = read_docx(file_path)
        elif lower_filename.endswith(".txt"):
            text = read_txt(file_path)
        
        if text.strip():
            documents.append((filename, text))
            logger.debug(f"已加载文档: {filename}, 字符数: {len(text)}")
    
    logger.info(f"共加载 {len(documents)} 个文档")
    return documents

def load_single_document(file_path: str) -> Optional[Tuple[str, str]]:
    """
    加载单个文档
    
    Args:
        file_path: 文件路径
        
    Returns:
        (文件名, 文件内容) 元组，如果加载失败返回 None
    """
    if not os.path.exists(file_path):
        logger.error(f"文件不存在: {file_path}")
        return None
    
    filename = os.path.basename(file_path)
    lower_filename = filename.lower()
    text = ""
    
    if lower_filename.endswith(".pdf"):
        text = read_pdf(file_path)
    elif lower_filename.endswith(".docx"):
        text = read_docx(file_path)
    elif lower_filename.endswith(".txt"):
        text = read_txt(file_path)
    else:
        logger.warning(f"不支持的文件格式: {filename}")
        return None
    
    if text.strip():
        return (filename, text)
    return None
